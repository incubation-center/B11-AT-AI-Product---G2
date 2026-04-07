import json
import logging
import io
import re
from datetime import datetime
from typing import Optional, Any

import pandas as pd
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.datasets import Dataset
from app.models.defects import Defect

logger = logging.getLogger(__name__)

# ─── Column mapping for CSV / Excel ─────────────────────────────────────────

DEFAULT_COLUMN_MAP = {
    "bug_id": [
        "bug_id", "bug id", "defect_id", "defect id", "id", "ticket", "ticket_id",
        "test_id", "test id", "testcase_id", "tc_id", "test case id", "case id",
        "case_id", "no", "no.", "number", "#",
    ],
    "title": [
        "title", "summary", "description", "defect_title", "bug_title", "name",
        "test_case", "test case", "scenario", "test_description", "test description",
        "test_name", "test name",
    ],
    "module": ["module", "component", "area", "feature", "module_name"],
    "severity": ["severity", "sev", "severity_level"],
    "priority": ["priority", "pri", "priority_level"],
    "environment": ["environment", "env", "platform", "test_environment"],
    "status": ["status", "state", "defect_status", "bug_status", "result"],
    "test_steps": ["test_steps", "test steps", "steps", "step_description", "step_desc"],
    "expected_result": ["expected_result", "expected result", "expected_outcome", "expected"],
    "preconditions": ["preconditions", "precondition", "precond", "setup"],
    "created_date": [
        "created_date", "created", "open_date", "opened", "date_created",
        "reported_date", "created_at",
    ],
    "resolved_date": ["resolved_date", "resolved", "fixed_date", "fix_date", "resolved_at"],
    "closed_date": ["closed_date", "closed", "close_date", "closed_at"],
}

SEVERITY_VALUES = {"Critical", "High", "Medium", "Low"}
PRIORITY_VALUES = {"P1", "P2", "P3", "P4"}
STATUS_VALUES = {"Open", "In Progress", "Passed", "Failed", "Blocked", "Automated", "Closed"}


def _normalize(col: str) -> str:
    return col.strip().lower().replace(" ", "_").replace("-", "_")


def _auto_map_columns(df_columns: list[str], explicit_mapping=None) -> dict[str, Optional[str]]:
    normalized = {_normalize(c): c for c in df_columns}
    mapping: dict[str, Optional[str]] = {}

    for db_field, aliases in DEFAULT_COLUMN_MAP.items():
        if explicit_mapping and getattr(explicit_mapping, db_field, None):
            explicit_val = getattr(explicit_mapping, db_field)
            if explicit_val in df_columns or _normalize(explicit_val) in normalized:
                mapping[db_field] = (
                    explicit_val if explicit_val in df_columns else normalized[_normalize(explicit_val)]
                )
                continue

        matched = None
        for alias in aliases:
            if alias in normalized:
                matched = normalized[alias]
                break
        mapping[db_field] = matched

    return mapping


def _parse_date(val) -> Optional[datetime]:
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return None
    if isinstance(val, datetime):
        return val.replace(tzinfo=None)
    if isinstance(val, pd.Timestamp):
        return val.to_pydatetime().replace(tzinfo=None)
    try:
        return pd.to_datetime(val, dayfirst=False).to_pydatetime().replace(tzinfo=None)
    except Exception:
        return None


def _safe_str(val) -> Optional[str]:
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return None
    return str(val).strip() or None


# ─── PDF / DOCX text extraction ──────────────────────────────────────────────

def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ValueError("PyMuPDF is not installed. Run: pip install pymupdf")
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n".join(pages)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a DOCX file using python-docx."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ValueError("python-docx is not installed. Run: pip install python-docx")
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            parts.append(row_text)
    return "\n".join(parts)


# ─── AI extraction for unstructured docs ─────────────────────────────────────

def _build_extraction_prompt(raw_text: str, file_name: str, max_cases: int = 100) -> str:
    # Limit raw text to avoid token explosion
    truncated = raw_text[:80_000] if len(raw_text) > 80_000 else raw_text
    return (
        "You are a QA engineering assistant. Extract ALL test cases from the document text below.\n"
        "Return ONLY valid JSON — a JSON array of objects. No markdown, no explanation.\n\n"
        f"Source file: {file_name}\n"
        f"Generate at most {max_cases} test cases.\n\n"
        "Each JSON object MUST have these keys:\n"
        "- bug_id (string): Test Case ID. If the document contains an ID (e.g. TC-001, TC001, QA-01), use it.\n"
        "  If the document has no ID, generate a sequential one: TC-001, TC-002, … etc.\n"
        "- title (string, required): Short action-oriented test case name.\n"
        "- module (string|null): Software module or feature being tested. Infer from context if not explicit.\n"
        "- severity (string|null): one of Critical, High, Medium, Low. Infer if not stated.\n"
        "- priority (string|null): one of P1, P2, P3, P4. Infer if not stated.\n"
        "- environment (string|null): test environment (e.g. staging, production, UAT). Infer if not stated.\n"
        "- status (string|null): one of Open, In Progress, Passed, Failed, Blocked, Automated. Use 'Open' if not stated.\n"
        "- preconditions (string|null): prerequisite conditions before the test is executed.\n"
        "- test_steps (string|null): Numbered step-by-step instructions for the tester. "
        "If missing, generate reasonable steps from context.\n"
        "- expected_result (string|null): The observable outcome after each step or at the end.\n\n"
        "RULES:\n"
        "- Every test case MUST have a bug_id — generate one if the document doesn't provide it.\n"
        "- title is required; never leave it empty.\n"
        "- Infer missing fields from context; only use null when you truly cannot determine the value.\n"
        "- Write in QA/tester language, not developer language.\n"
        "- Each test case should be self-contained and independently executable.\n\n"
        "=== DOCUMENT TEXT ===\n"
        f"{truncated}\n"
        "=== END DOCUMENT TEXT ==="
    )


def _extract_json_array(text: str) -> list[Any]:
    text = text.strip()
    try:
        parsed = json.loads(text)
        if isinstance(parsed, list):
            return parsed
    except Exception:
        pass
    start = text.find("[")
    end = text.rfind("]")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("AI output did not contain a JSON array")
    parsed = json.loads(text[start: end + 1])
    if not isinstance(parsed, list):
        raise ValueError("Extracted JSON was not an array")
    return parsed


async def _ai_extract_testcases(raw_text: str, file_name: str) -> list[dict[str, Any]]:
    """Use the LLM (via OpenRouter) to extract structured test cases from raw document text."""
    from app.services.embedding_service import _generate_with_retry, GENERATIVE_MODEL

    prompt = _build_extraction_prompt(raw_text, file_name)
    output = await _generate_with_retry(GENERATIVE_MODEL, prompt)
    items = _extract_json_array(output)
    return [item for item in items if isinstance(item, dict)]


def _normalize_ai_item(item: dict[str, Any], index: int) -> dict[str, Any]:
    """Post-process a single AI-extracted test case, filling blanks and normalising values."""

    def opt(key: str) -> Optional[str]:
        v = item.get(key)
        if v is None:
            return None
        v = str(v).strip()
        return v if v else None

    # Ensure bug_id always has a value
    bug_id = opt("bug_id") or f"TC-{index + 1:03d}"

    # Normalise severity / priority to known values
    severity = opt("severity")
    if severity and severity.capitalize() in SEVERITY_VALUES:
        severity = severity.capitalize()
    elif severity and severity.upper() in SEVERITY_VALUES:
        severity = severity.upper()

    priority = opt("priority")
    if priority and priority.upper() in PRIORITY_VALUES:
        priority = priority.upper()

    status = opt("status") or "Open"

    return {
        "bug_id": bug_id,
        "title": opt("title") or "Untitled Test Case",
        "module": opt("module"),
        "severity": severity,
        "priority": priority,
        "environment": opt("environment"),
        "status": status,
        "preconditions": opt("preconditions"),
        "test_steps": opt("test_steps"),
        "expected_result": opt("expected_result"),
    }


# ─── Main entry point ─────────────────────────────────────────────────────────

async def _process_with_ai(
    db: AsyncSession,
    user_id: int,
    file_name: str,
    raw_text: str,
    stored_type: str,
    upload_type: str,
) -> tuple[Dataset, int]:
    if not raw_text.strip():
        raise ValueError("The document content appears to be empty or could not be read.")

    logger.info(f"Running AI extraction on {len(raw_text[:10000])} chars (truncated metric) from {file_name}…")
    ai_items = await _ai_extract_testcases(raw_text, file_name)

    if not ai_items:
        raise ValueError("No test cases could be extracted from the document.")

    dataset = Dataset(
        user_id=user_id,
        file_name=file_name,
        file_type=stored_type,
        upload_type=upload_type,
    )
    db.add(dataset)
    await db.flush()

    defects: list[Defect] = []
    for i, item in enumerate(ai_items):
        norm = _normalize_ai_item(item, i)
        defects.append(Defect(
            dataset_id=dataset.dataset_id,
            bug_id=norm["bug_id"],
            title=norm["title"],
            module=norm["module"],
            severity=norm["severity"],
            priority=norm["priority"],
            environment=norm["environment"],
            status=norm["status"],
            preconditions=norm["preconditions"],
            test_steps=norm["test_steps"],
            expected_result=norm["expected_result"],
        ))

    db.add_all(defects)
    await db.flush()
    logger.info(f"AI-extracted {len(defects)} test cases from '{file_name}'")
    return dataset, len(defects)


async def parse_and_import(
    db: AsyncSession,
    user_id: int,
    file_name: str,
    file_bytes: bytes,
    file_type: str,
    upload_type: str,
    explicit_mapping=None,
) -> tuple[Dataset, int]:
    """
    Parse an uploaded file (CSV, Excel, PDF, DOCX) and bulk-insert Defects.
    - CSV / Excel: column-mapped structured extraction.
    - PDF / DOCX / DOC: AI-powered unstructured extraction with auto-fill.
    Returns (dataset, defect_count).
    """
    file_type_lower = file_type.lower()

    # ── Structured: CSV / Excel ───────────────────────────────────────────────
    if file_type_lower in ("csv", "text/csv", "application/csv"):
        df = pd.read_csv(io.BytesIO(file_bytes))
        stored_type = "csv"

    elif file_type_lower in (
        "xlsx", "xls",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        df = pd.read_excel(io.BytesIO(file_bytes), engine="openpyxl")
        stored_type = "excel"

    else:
        # ── Unstructured: PDF / DOCX / DOC ───────────────────────────────────
        if file_type_lower in ("pdf", "application/pdf"):
            raw_text = _extract_text_from_pdf(file_bytes)
            stored_type = "pdf"
        elif file_type_lower in (
            "docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "doc",
            "application/msword",
        ):
            raw_text = _extract_text_from_docx(file_bytes)
            stored_type = "docx"
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        return await _process_with_ai(db, user_id, file_name, raw_text, stored_type, upload_type)

    # ── Structured path continues here (CSV / Excel) ──────────────────────────
    if df.empty:
        raise ValueError("The uploaded file contains no data rows.")

    logger.info(f"Parsed {len(df)} rows, columns: {list(df.columns)}")

    col_map = _auto_map_columns(list(df.columns), explicit_mapping)
    logger.info(f"Column mapping: {col_map}")

    if col_map.get("title") is None:
        logger.warning(
            f"Could not detect a 'title' / 'summary' column in '{file_name}'. "
            f"Falling back to AI extraction. Available columns: {list(df.columns)}"
        )
        raw_text = df.to_csv(index=False)
        return await _process_with_ai(db, user_id, file_name, raw_text, stored_type, upload_type)

    dataset = Dataset(
        user_id=user_id,
        file_name=file_name,
        file_type=stored_type,
        upload_type=upload_type,
    )
    db.add(dataset)
    await db.flush()

    defects: list[Defect] = []
    for idx, (_, row) in enumerate(df.iterrows()):
        # Auto-generate bug_id if missing
        raw_bug_id = _safe_str(row.get(col_map["bug_id"])) if col_map.get("bug_id") else None
        bug_id = raw_bug_id or f"TC-{idx + 1:03d}"

        defect = Defect(
            dataset_id=dataset.dataset_id,
            bug_id=bug_id,
            title=_safe_str(row[col_map["title"]]) or "Untitled",
            module=_safe_str(row.get(col_map["module"])) if col_map.get("module") else None,
            severity=_safe_str(row.get(col_map["severity"])) if col_map.get("severity") else None,
            priority=_safe_str(row.get(col_map["priority"])) if col_map.get("priority") else None,
            environment=_safe_str(row.get(col_map["environment"])) if col_map.get("environment") else None,
            status=_safe_str(row.get(col_map["status"])) if col_map.get("status") else None,
            test_steps=_safe_str(row.get(col_map["test_steps"])) if col_map.get("test_steps") else None,
            expected_result=_safe_str(row.get(col_map["expected_result"])) if col_map.get("expected_result") else None,
            preconditions=_safe_str(row.get(col_map["preconditions"])) if col_map.get("preconditions") else None,
            created_date=_parse_date(row.get(col_map["created_date"])) if col_map.get("created_date") else None,
            resolved_date=_parse_date(row.get(col_map["resolved_date"])) if col_map.get("resolved_date") else None,
            closed_date=_parse_date(row.get(col_map["closed_date"])) if col_map.get("closed_date") else None,
        )
        defects.append(defect)

    db.add_all(defects)
    await db.flush()

    return dataset, len(defects)
